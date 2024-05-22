import io,time
import os
from flask import Flask, request, render_template, send_file, send_from_directory, session, redirect, url_for
from flask_socketio import SocketIO, emit
from docx import Document
import logging
from docx.enum.text import WD_COLOR_INDEX


app = Flask(__name__)
socketio = SocketIO(app)
app.secret_key = 'supersecretkey'  # Replace with a real secret key
socketio = SocketIO(app)
logging.basicConfig(level=logging.DEBUG)

DOCUMENTS = {
    'rental': 'residential-rental-agreement-format.docx',
    'nda': 'residential-rental-agreement-format.docx',
    'poa': 'residential-rental-agreement-format.docx'
}

@app.route('/')
def home():
    return render_template('index.html', documents=DOCUMENTS)

@app.route('/form', methods=['GET'])
def show_form():
    doc_type = request.args.get('doc_type')
    document_path = DOCUMENTS.get(doc_type)
    if not document_path:
        return 'Document not found.', 404

    document = Document(document_path)
    placeholders = extract_placeholders(document)
    return render_template('generate.html', placeholders=placeholders, doc_type=doc_type)

def extract_placeholders(document):
    placeholders = []
    for paragraph in document.paragraphs:
        if '[' in paragraph.text and ']' in paragraph.text:
            placeholders.extend(extract_from_text(paragraph.text))
    return placeholders

def extract_from_text(text):
    return [part[1:-1] for part in text.split('[') if ']' in part]

def replace_text(doc, replacements):
    for para in doc.paragraphs:
        replace_in_paragraph(para, replacements)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    replace_in_paragraph(para, replacements)

def replace_in_paragraph(paragraph, replacements):
    for key, value in replacements.items():
        if key in paragraph.text:
            paragraph.text = paragraph.text.replace(key, value)

@app.route('/generate', methods=['POST'])
def generate_document():
    print("Received form data:", request.form)  # Debugging line to print received form data

    doc_type = request.form.get('doc_type')
    document_path = DOCUMENTS.get(doc_type, 'residential-rental-agreement-format.docx')
    doc = Document(document_path)

    replacements = {
        '[City, State]': request.form.get('CityState', '[City, State]'),
        '[Date, Month, Year]': request.form.get('DateMonthYear', '[Date, Month, Year]'),
        '[Landlord Name]': request.form.get('LandlordName', '[Landlord Name]'),
        '[Landlord Address Line 1, Address Line 2, City, State, Pin Code]': request.form.get('LandlordAddress', '[Landlord Address Line 1, Address Line 2, City, State, Pin Code]'),
        '[Tenant Name]': request.form.get('TenantName', '[Tenant Name]'),
        '[Tenant Address Line 1, Address Line 2, City, State, Pin Code]': request.form.get('TenantAddress', '[Tenant Address Line 1, Address Line 2, City, State, Pin Code]'),
        '[Lease Property Address Line 1, Address Line 2, City, State, Pin Code]': request.form.get('PropertyAddress', '[Lease Property Address Line 1, Address Line 2, City, State, Pin Code]'),
        '[Independent House / Apartment / Farm House / Residential Property]': request.form.get('PropertyCategory', '[Independent House / Apartment / Farm House / Residential Property]'),
        '[X Bedrooms], [X Bathrooms], [X Carparks]': request.form.get('PropertyDetails', '[X Bedrooms], [X Bathrooms], [X Carparks]'),
        '[Lease Term]': request.form.get('LeaseTerm', '[Lease Term]'),
        '[Lease Start Date]': request.form.get('LeaseStartDate', '[Lease Start Date]'),
        '[Monthly Rental in Number & Words]': request.form.get('MonthlyRental', '[Monthly Rental in Number & Words]'),
        '[Rental Deposit in Number and Words]': request.form.get('RentalDeposit', '[Rental Deposit in Number and Words]'),
        '[Starting Meter Reading]': request.form.get('StartMeterReading', '[Starting Meter Reading]'),
        '[XXXX Square Feet]': request.form.get('SquareFeet', '[XXXX Square Feet]'),
        '[Lessor Name]': request.form.get('LessorName', '[Lessor Name]'),
        '[Lessor Address Line 1]': request.form.get('LessorAddress1', '[Lessor Address Line 1]'),
        '[Lessor Address Line 2]': request.form.get('LessorAddress2', '[Lessor Address Line 2]'),
        '[Lessor City, State, Pin Code]': request.form.get('LessorCityStatePin', '[Lessor City, State, Pin Code]'),
        '[Lessee Name]': request.form.get('LesseeName', '[Lessee Name]'),
        '[Lessee Address Line 1]': request.form.get('LesseeAddress1', '[Lessee Address Line 1]'),
        '[Lessee Address Line 2]': request.form.get('LesseeAddress2', '[Lessee Address Line 2]'),
        '[Lessee City, State, Pin Code]': request.form.get('LesseeCityStatePin', '[Lessee City, State, Pin Code]'),
        '[WITNESS ONE Name & Address]': request.form.get('WitnessOneNameAddress', '[WITNESS ONE Name & Address]'),
        '[WITNESS TWO Name & Address]': request.form.get('WitnessTwoNameAddress', '[WITNESS TWO Name & Address]')
    }

    replace_text(doc, replacements)

    # Save the modified document locally
    output_filename = 'Generated_Document.docx'
    doc.save(output_filename)

    # Store only the filename in session
    session['output_filename'] = output_filename

    # Redirect to a new route that will handle the streaming
    return redirect(url_for('stream_content'))

@app.route('/stream_content')
def stream_content():
    output_filename = session.get('output_filename')
    if not output_filename:
        return "No content available.", 404

    return render_template('generating_document.html')

@socketio.on('stream_request')
def handle_stream_request(json, methods=['GET', 'POST']):
    print('Received stream request: ' + str(json))
    output_filename = session.get('output_filename')
    if not output_filename:
        emit('stream_response', {'word': 'No content available.'})
        return

    doc_path = os.path.join(os.getcwd(), output_filename)
    doc = Document(doc_path)
    full_text = get_full_text(doc)
    words = full_text.split()
    
    for word in words:
        emit('stream_response', {'word': word})
        time.sleep(0.1)  # Slight delay to simulate streaming

def get_full_text(doc):
    full_text = []
    for para in doc.paragraphs:
        full_text.append(para.text)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for para in cell.paragraphs:
                    full_text.append(para.text)
    return '\n'.join(full_text)

@app.route('/download_document')
def download_document():
    output_filename = session.get('output_filename', 'Generated_Document.docx')
    return send_from_directory(directory=os.getcwd(), path=output_filename, as_attachment=True)
@app.route('/contact')
def contact():
    return render_template('contact.html')

